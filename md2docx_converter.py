
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_to_docx(input_file, output_file):
    doc = Document()
    
    # Style settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_buffer = []
    
    for line in lines:
        stripped = line.strip()
        
        # Handle Tables
        if stripped.startswith('|'):
            table_buffer.append(stripped)
            continue
        else:
            if table_buffer:
                # Process collected table
                process_table(doc, table_buffer)
                table_buffer = []

        if not stripped:
            continue
            
        # Headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            
        # Horizontal Rule
        elif stripped.startswith('---'):
            p = doc.add_paragraph()
            run = p.add_run('________________________________________________________________')
            run.font.color.rgb = RGBColor(200, 200, 200)
            
        # Lists
        elif stripped.startswith('* ') or stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
            
        # Blockquotes (Simple fallback)
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            
        # Normal Text
        else:
            # Simple bold parsing (**text**)
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Flush remaining table if file ends with table
    if table_buffer:
        process_table(doc, table_buffer)

    doc.save(output_file)
    print(f"Successfully converted {input_file} to {output_file}")

def process_table(doc, lines):
    # Determine columns from the separator line (usually 2nd line)
    # Filter out the separator line |---|---|
    data_lines = [l for l in lines if not set(l.strip()).issubset({'|', '-', ':', ' '})]
    
    if not data_lines:
        return

    # Parse headers (first line)
    header_raw = data_lines[0].strip().strip('|').split('|')
    headers = [h.strip() for h in header_raw]
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Parse body
    for line in data_lines[1:]:
        row_raw = line.strip().strip('|').split('|')
        row_data = [d.strip() for d in row_raw]
        
        row_cells = table.add_row().cells
        for i, datum in enumerate(row_data):
            if i < len(row_cells):
                 row_cells[i].text = datum

if __name__ == "__main__":
    parse_markdown_to_docx("DOCUMENTACAO_TECNICA.md", "PolitiRank_Documentacao_Tecnica.docx")
